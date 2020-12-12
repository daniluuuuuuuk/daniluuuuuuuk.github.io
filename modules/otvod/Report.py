from qgis.core import QgsProject, QgsFeature, QgsMapSettings, QgsMapRendererSequentialJob, QgsCoordinateReferenceSystem
from qgis.PyQt.QtWidgets import QMessageBox
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from qgis.PyQt.QtCore import QSize
from PIL import Image, ImageOps
from ...tools import config

class Report:

    def __init__(self, table, canvas, layers):
        self.table = table
        self.canvas = canvas
        self.uid = self.getUid()
        self.feature = self.getFeatureByUid()
        self.areaAttributes = self.getAreaAttributes()
        self.otvodAttributes = self.getOtvodAttributes()
        self.columnNames = self.extractColumnNames(self.table.getColumnNames())
        self.layers = layers

    def generate(self):

        self.saveImg(self.layers)

        document = Document()

        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        header = document.add_paragraph()
        header.add_run('КАРТА-СХЕМА').bold = True
        header.add_run(
            '\nс обозначенными границами участка лесного фонда, предоставляемого для лесопользования\n').bold = True
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(
            'Месторасположение лесосеки:\n'
            'Юридическое лицо, ведущее лесное хозяйство:'
        )
        p.paragraph_format.left_indent = Inches(0.25)

        p = document.add_paragraph(
            'Структурное подразделение юридического лица, ведущего лесное хозяйство:_______________'
            ', лесной квартал №{}, таксационный выдел №{}, площадь лесосеки {} га.'.format(str(
                self.areaAttributes["num_kv"]), str(self.areaAttributes["num_vds"]), str(self.areaAttributes["area"]))
        )
        p.paragraph_format.first_line_indent = Inches(0.25)

        p = document.add_paragraph('Масштаб: 1:10000')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0.25)

        in_img = os.path.join(QgsProject.instance().homePath(), 'output.png')
        self.add_border(in_img, output_image=os.path.join(
            QgsProject.instance().homePath(), "outputborder.png"), border=1)

        document.add_picture(os.path.join(
            QgsProject.instance().homePath(), "outputborder.png"))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('Условные обозначения:')

        p = document.add_paragraph(
            'Экспликация или координаты поворотных точек лесосеки:')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=1, cols=len(self.columnNames))
        table.style = 'TableGrid'

        hdr_cells = table.rows[0].cells

        for x in range(0, len(self.columnNames)):
            if self.columnNames[x] == 'GPS':
                pass
            else:
                hdr_cells[x].text = self.columnNames[x]
                hdr_cells[x].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(0, len(self.otvodAttributes)):
            row_cells = table.add_row().cells
            for y in range(len(self.otvodAttributes[i])):
                row_cells[y].text = str(self.otvodAttributes[i][y])
                row_cells[y].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(
            '\nИсполнитель:{}'.format(self.areaAttributes["fio"]))

        path = os.path.join(self.getReportFolder(), 'Отвод_кв_{}.docx'.format(str(self.areaAttributes["num_kv"])))
        document.save(path)
        path = path.replace("\\", "\\\\")
        path = path.replace("/", "\\\\")
        path = path.replace(" ", "%20")

        return path

    def getReportFolder(self):
        cf = config.Configurer('report')
        settings = cf.readConfigs()
        return settings.get('path')

    def extractColumnNames(self, columnNames):
        newColumns = []
        for col in columnNames:
            if col == 'GPS':
                pass
            else:
                newColumns.append(col)
        return newColumns

    def getOtvodAttributes(self):
        otvodDict = []
        data = self.table.getJSONRows()
        for row in data[1:]:
            for parameter in row:
                dictt = []
                for key, value in (row[parameter].items()):
                    dictt.append(value)
            otvodDict.append(dictt)
        return otvodDict

    def getUid(self):
        try:
            layer = QgsProject.instance().mapLayersByName(
                "Лесосека временный слой")[0]
            features = []
            for feature in layer.getFeatures():
                features.append(feature)
            return features[0]['uid']
        except Exception as e:
            QMessageBox.information(None, "Ошибка модуля QGis", str(e))

    def getFeatureByUid(self):
        try:
            layer = QgsProject.instance().mapLayersByName("Лесосеки")[0]
            for feature in layer.getFeatures():
                if feature['uid'] == self.uid:
                    return feature
        except:
            QMessageBox.information(
                None, "Ошибка модуля QGis", "Лесосека с таким идентификатором не найдена " + str(self.uid))

    def getAreaAttributes(self):
        try:
            attrDict = {}
            attrDict["num_kv"] = self.feature["num_kv"]
            attrDict["num_vds"] = self.feature["num_vds"]
            attrDict["area"] = self.feature["area"]
            attrDict["num"] = self.feature["num"]
            attrDict["useType"] = self.feature["usetype"]
            attrDict["cuttingType"] = self.feature["cuttingtyp"]
            attrDict["plot"] = self.feature["plot"]
            attrDict["fio"] = self.feature["fio"]
            attrDict["date"] = self.feature["date"]
            attrDict["info"] = self.feature["info"]
            return attrDict
        except Exception as e:
            QMessageBox.information(
                None, "Ошибка модуля QGis", "Сначала сохраните лесосеку " + str(e))

    def saveImg(self, layers):
        settings = QgsMapSettings()
        # print('На входе', settings.scale(), self.canvas.scale())
        # print('Первый скейл', settings.scale(), self.canvas.scale())
        self.canvas.refresh()
        # print('После рефреша', settings.scale(), self.canvas.scale())
        settings.setOutputSize(QSize(391, 361))
        self.canvas.zoomScale(10000)
        self.canvas.refresh()
        # print('После аутпут сайза', settings.scale(), self.canvas.scale())
        settings.setExtent(self.canvas.extent())
        # print('После сет экстента', settings.scale(), self.canvas.scale())
        # settings.setOutputDpi(126)
        # print('После дпи', settings.scale(), self.canvas.scale())
        settings.setLayers(layers)
        # settings.setEllipsoid("WGS84")
        # print('После эллипсоида', settings.scale(), self.canvas.scale())
        crs = QgsCoordinateReferenceSystem('EPSG:32635')
        # print('После координат', settings.scale(), self.canvas.scale())
        settings.setDestinationCrs(crs)
        # print('После дестинейшен СРС', settings.scale(), self.canvas.scale())
        # self.canvas.zoomScale(10000)
        # self.canvas.refresh()
        job = QgsMapRendererSequentialJob(settings)
        job.start()
        job.waitForFinished()
        img = job.renderedImage()
        path2 = os.path.join(QgsProject.instance().homePath(), "output.png")
        img.save(path2)
        # print('В конце', settings.scale(), self.canvas.scale(), settings.mapUnits())

    def add_border(self, input_image, output_image, border):
        img = Image.open(input_image)
        if isinstance(border, int) or isinstance(border, tuple):
            bimg = ImageOps.expand(img, border=border, fill='black')
        else:
            raise RuntimeError('Border is not an image or tuple')
        bimg.save(output_image, dpi=(96, 96))
